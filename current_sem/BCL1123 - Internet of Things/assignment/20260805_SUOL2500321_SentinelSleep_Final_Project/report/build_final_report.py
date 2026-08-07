from __future__ import annotations

import html
import re
from pathlib import Path

import markdown
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "20260805_SUOL2500321_SentinelSleep_Final_Report.md"
OUT_DOCX = ROOT / "20260805_SUOL2500321_SentinelSleep_Final_Report.docx"
OUT_HTML = ROOT / "20260805_SUOL2500321_SentinelSleep_Final_Report.html"
ASSETS = ROOT / "assets"
TNR = "Times New Roman"


def set_cell_margins(cell, value=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Chan Jing Yi | SUOL2500321 | Page ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def configure_section(section, landscape=False):
    section.page_width = Mm(297 if landscape else 210)
    section.page_height = Mm(210 if landscape else 297)
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    header = section.header.paragraphs[0]
    header.text = "BCL1123 Internet of Things | SentinelSleep Final Project"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.name = TNR
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)
    add_page_number(section.footer.paragraphs[0])
    for run in section.footer.paragraphs[0].runs:
        run.font.name = TNR
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)


def add_toc(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click and update this field if page numbers are not visible."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_inline(paragraph, source, size=12, bold=False, italic=False):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)", source)
    for part in parts:
        if not part:
            continue
        code = part.startswith("`") and part.endswith("`")
        strong = part.startswith("**") and part.endswith("**")
        emph = part.startswith("*") and part.endswith("*") and not strong
        text = part[1:-1] if code or emph else part[2:-2] if strong else part
        run = paragraph.add_run(text)
        run.font.name = "Courier New" if code else TNR
        run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
        run.font.size = Pt(10 if code else size)
        run.bold = bold or strong
        run.italic = italic or emph


def add_body_paragraph(document, text, note=False):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Inches(0 if note else 0.3)
    if note:
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.right_indent = Inches(0.3)
    add_inline(paragraph, text, bold=note)
    return paragraph


def add_image(document, relative_path, alt_text):
    path = ROOT / relative_path
    if not path.exists():
        path.parent.mkdir(parents=True, exist_ok=True)
        placeholder = Image.new("RGB", (1280, 720), color=(15, 23, 42))
        draw = ImageDraw.Draw(placeholder)
        draw.rectangle([(20, 20), (1260, 700)], outline=(59, 130, 246), width=4)
        draw.text((640, 360), f"Visual Evidence: {alt_text}", fill=(248, 250, 252), anchor="mm")
        placeholder.save(path)
    with Image.open(path) as img:
        width, height = img.size
    section = document.sections[-1]
    content_width = section.page_width - section.left_margin - section.right_margin
    max_inches = min(content_width / 914400, 9.4 if section.orientation == WD_ORIENT.LANDSCAPE else 6.45)
    max_height = 6.8 if section.orientation == WD_ORIENT.PORTRAIT else 5.4
    target_width = max_inches
    target_height = target_width * height / width
    if target_height > max_height:
        target_height = max_height
        target_width = target_height * width / height
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(target_width), height=Inches(target_height))
    drawing = run._r.xpath(".//wp:docPr")
    if drawing:
        drawing[0].set("descr", alt_text)
        drawing[0].set("title", alt_text)


def add_table(document, rows):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline(paragraph, value, size=9.5, bold=r_idx == 0)
            if r_idx == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "DDEBF2")
                cell._tc.get_or_add_tcPr().append(shading)
    document.add_paragraph()


def add_placeholder_assets():
    font_path = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
    small_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    for filename, title, body in [
        ("figure3_wokwi_pending.png", "WOKWI RUNTIME EVIDENCE PENDING", "Capture the full circuit and serial monitor after Blynk authentication."),
        ("figure6_blynk_pending.png", "BLYNK DASHBOARD EVIDENCE PENDING", "Capture the live dashboard with readings, trends, controls and online state."),
    ]:
        out = ASSETS / filename
        if out.exists():
            continue
        image = Image.new("RGB", (1600, 900), "#F4F7FA")
        draw = ImageDraw.Draw(image)
        draw.rounded_rectangle((80, 80, 1520, 820), radius=30, outline="#B42318", width=7, fill="#FFF7F6")
        title_font = ImageFont.truetype(font_path, 48)
        body_font = ImageFont.truetype(small_path, 32)
        draw.text((800, 360), title, font=title_font, fill="#B42318", anchor="mm")
        draw.text((800, 450), body, font=body_font, fill="#34495E", anchor="mm")
        image.save(out)


def build_docx():
    add_placeholder_assets()
    document = Document()
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = TNR
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    for name, size, color in [("Title", 22, "17324D"), ("Heading 1", 16, "17324D"), ("Heading 2", 13, "0B6E99")]:
        style = styles[name]
        style.font.name = TNR
        style._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    configure_section(document.sections[0], landscape=False)
    cover = document.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(100)
    add_inline(cover, "SEGi UNIVERSITY", size=16, bold=True)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    add_inline(title, "SENTINELSLEEP\nSMART BEDROOM IOT SYSTEM", size=22, bold=True)
    for text in [
        "BCL1123 Internet of Things",
        "Final Report & Video",
        "May–August 2026",
        "Chan Jing Yi",
        "SUOL2500321",
        "Degree IoT ODL",
        "Lecturer: Lee Thian Seng",
        "Submission date: As in LMS",
    ]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(8)
        add_inline(paragraph, text, size=12, bold=text in {"Chan Jing Yi", "SUOL2500321"})
    signature = ASSETS / "signature_no_bg.png"
    if signature.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(signature), width=Inches(1.3))
        caption = document.add_paragraph("Student signature")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    document.add_heading("Table of Contents", level=1)
    add_toc(document)
    document.add_page_break()

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1."))
    i = start
    in_landscape = False
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("### 3.3") and not in_landscape:
            section = document.add_section(WD_SECTION.NEW_PAGE)
            configure_section(section, landscape=True)
            in_landscape = True
        elif line.startswith("### 3.5") and in_landscape:
            section = document.add_section(WD_SECTION.NEW_PAGE)
            configure_section(section, landscape=False)
            in_landscape = False
        if line.startswith("## "):
            if i != start:
                document.add_page_break()
            document.add_heading(line[3:], level=1)
            i += 1
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            i += 1
            continue
        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", line)
        if image_match:
            add_image(document, image_match.group(2), image_match.group(1))
            i += 1
            continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    rows.append(cells)
                i += 1
            add_table(document, rows)
            continue
        if line.startswith("```"):
            language = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            i += 1
            continue
        if line.startswith("**Figure"):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            add_inline(paragraph, line, size=10, italic=True)
            i += 1
            continue
        if line.startswith("> "):
            note_text = line[2:]
            add_body_paragraph(document, note_text, note=True)
            if "Figure 3" in note_text:
                add_image(document, "assets/figure3_wokwi_pending.png", "Wokwi runtime evidence placeholder")
            elif "Figure 6" in note_text:
                add_image(document, "assets/figure6_blynk_pending.png", "Blynk dashboard evidence placeholder")
            i += 1
            continue
        add_body_paragraph(document, line)
        i += 1

    settings = document.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)
    document.core_properties.title = "SentinelSleep Smart Bedroom IoT System — Final Project Report"
    document.core_properties.author = "Chan Jing Yi"
    document.core_properties.subject = "BCL1123 Internet of Things Final Project"
    document.core_properties.keywords = "IoT, ESP32, Wokwi, Blynk, smart bedroom"
    document.save(OUT_DOCX)


def build_html():
    add_placeholder_assets()
    source = SOURCE.read_text(encoding="utf-8")
    source = source.replace(
        "> Insert Figure 3 after the first authenticated Wokwi run: capture the full circuit, active serial monitor and visible output indicators without exposing the Blynk token.",
        "> Insert Figure 3 after the first authenticated Wokwi run: capture the full circuit, active serial monitor and visible output indicators without exposing the Blynk token.\n\n![Wokwi runtime evidence pending](assets/figure3_wokwi_pending.png)",
    )
    source = source.replace(
        "> Insert Figure 6 after account setup: show the live SentinelSleep Blynk dashboard with current readings, trend chart, confirmed output indicators, controls and device-online status.",
        "> Insert Figure 6 after account setup: show the live SentinelSleep Blynk dashboard with current readings, trend chart, confirmed output indicators, controls and device-online status.\n\n![Blynk dashboard evidence pending](assets/figure6_blynk_pending.png)",
    )
    body = markdown.markdown(source, extensions=["tables", "fenced_code", "toc"], output_format="html5")
    page = f"""<!doctype html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>SentinelSleep Final Project Report</title>
<style>
@page {{ size:A4; margin:20mm; }}
body {{ max-width:900px; margin:0 auto; font-family:'Times New Roman',serif; font-size:12pt; line-height:1.5; color:#182a38; }}
h1 {{ text-align:center; color:#17324d; }} h2 {{ color:#17324d; page-break-before:always; border-bottom:1px solid #aac1d0; }} h3 {{ color:#0b6e99; page-break-after:avoid; }}
p {{ text-align:justify; text-indent:2em; }} strong {{ color:#17324d; }}
blockquote {{ border-left:5px solid #b42318; background:#fff4f2; padding:10px 16px; margin:14px 0; }} blockquote p {{ text-indent:0; font-weight:bold; }}
table {{ width:100%; border-collapse:collapse; margin:12px 0; font-size:9.5pt; page-break-inside:auto; }} tr {{ page-break-inside:avoid; }} th {{ background:#ddebf2; }} th,td {{ border:1px solid #879eac; padding:6px; vertical-align:top; }}
img {{ display:block; max-width:100%; max-height:680px; margin:12px auto; object-fit:contain; page-break-inside:avoid; }}
pre {{ background:#f4f6f8; border:1px solid #c8d2da; padding:10px; white-space:pre-wrap; font-size:8.5pt; page-break-inside:avoid; }}
a {{ color:#0b6e99; overflow-wrap:anywhere; }}
body > h1:first-of-type {{ margin-top:30mm; font-size:24pt; }}
body > h2:first-of-type {{ page-break-before:avoid; text-align:center; border:0; }}
@media print {{ body {{ max-width:none; }} }}
</style></head><body>{body}</body></html>"""
    OUT_HTML.write_text(page, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_html()
    print(OUT_DOCX)
    print(OUT_HTML)

