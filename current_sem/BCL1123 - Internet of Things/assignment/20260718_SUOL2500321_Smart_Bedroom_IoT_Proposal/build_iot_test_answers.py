from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
SOURCE = BASE_DIR / "renametoyourstudentID 2.docx"
OUTPUT = BASE_DIR / "20260801_SUOL2500321_BCL1123_IoT_Test_Answered.docx"


ANSWERS = {
    87: [
        "The M2M era centred on direct, task-specific communication between a small number of machines. A sensor or controller normally sent telemetry to another machine through a proprietary wired link, cellular modem, or closed supervisory system. Human involvement was limited, data often remained inside an organisational silo, and adding a new vendor’s device usually required a custom gateway. M2M therefore prioritised reliable point-to-point automation rather than broad service integration.",
        "The modern IoT ecosystem connects physical and virtual things through Internet Protocol networks, cloud or edge platforms, application programming interfaces, and user-facing applications. A single device’s data can be stored, analysed, combined with other sources, and used by several services. IoT also supports many-to-many communication, remote device management, scalable analytics, and human interaction through dashboards or mobile alerts. This openness creates more value than isolated M2M telemetry, although it also introduces wider cybersecurity, privacy, governance, and interoperability responsibilities.",
    ],
    97: [
        "The Middleware or Processing Layer sits between device connectivity and business applications. Devices made by different manufacturers may transmit different payload formats and use protocols such as MQTT, HTTP, or CoAP. Middleware hides that heterogeneity through protocol adapters, gateways, device registries, and common application programming interfaces. It identifies each device, translates messages into a shared data model, validates units and timestamps, and exposes normalised information to applications. A dashboard can therefore consume one consistent temperature field even when the original sensors encode that reading differently.",
        "The layer also performs message brokering, event processing, data filtering, aggregation, storage, and rule execution. For example, it can convert Fahrenheit to Celsius, reject an impossible reading, combine several sensor values, and publish a standard alarm event. Semantic metadata clarifies what a value represents, while service discovery enables authorised applications to find the appropriate device capability without depending on its manufacturer’s implementation.",
        "Interoperability still requires governance. Middleware should enforce standard schemas and APIs, maintain version compatibility, authenticate devices, authorise access, encrypt data in transit, and record audit logs. It does not automatically correct a poorly documented proprietary protocol; an adapter or vendor-supported standard is still required. With those controls, applications remain decoupled from hardware brands and devices can be replaced or expanded without redesigning the complete IoT service.",
    ],
    111: [
        "The first critical component is a calibrated multi-sensor data logger placed with the biological samples. A digital temperature probe provides a continuous, time-stamped record, while a humidity sensor and door-open or shock sensor can reveal condensation, handling, or package breaches. The logger needs local memory and a backup battery so that evidence is retained during aircraft loading, customs inspection, or a temporary loss of network coverage. Calibration records and device identity link every reading to the correct shipment.",
        "The second component is a secure communications and location unit, consisting of a low-power controller, GNSS positioning, and cellular or LPWAN connectivity. It periodically sends the sensor readings, location, battery state, and shipment identifier to an IoT platform. Edge rules should trigger an audible or visual local warning when a safe threshold is crossed, while the platform alerts the control centre. Buffered store-and-forward transmission prevents a coverage gap from becoming a gap in the compliance record.",
    ],
    123: [
        "Real-time analytics converts the sensor stream into operational decisions. The platform compares each reading with the product-specific temperature and humidity limits, checks the duration of an excursion, and combines the result with location, route, weather, and estimated arrival time. A predictive rule can warn staff before the sample becomes unusable—for example, by detecting a rising temperature trend rather than waiting for a fixed limit to be exceeded. The control centre can then instruct a driver to inspect the container, increase cooling, move the shipment to qualified storage, or reroute it to a nearer approved facility.",
        "That early intervention reduces waste because the firm does not automatically discard every shipment after a minor event. A complete time–temperature history supports a risk-based decision on whether the biological sample remained within its validated stability envelope. Analytics can also compare lanes, carriers, containers, and packaging designs to identify repeated delays or thermal weak points. Maintenance can then be scheduled before a refrigeration unit fails, and inefficient routes or excessive refrigerant use can be corrected.",
        "Compliance improves when the platform preserves calibrated, time-stamped, tamper-evident records and applies the correct rules for each product and destination. Automated exception reports, chain-of-custody logs, acknowledgement records, and audit trails give regulators and customers consistent evidence. Role-based access and encryption protect sensitive shipment data. Sensor calibration drift, false alarms, and missing connectivity remain risks, so the firm should use calibration schedules, redundant critical sensors, local buffering, and human review before releasing or rejecting a shipment.",
    ],
    149: [
        "The first sensor would be a passive infrared (PIR) occupancy sensor positioned to detect movement around the bedroom study desk. It is appropriate because the main requirement is to distinguish an occupied work area from an empty room without recording identifiable video. Its low data rate and modest power requirement also suit continuous indoor monitoring. A short timeout can reduce false “vacant” decisions when the user is sitting relatively still.",
        "The second sensor would be a digital ambient-light sensor such as the BH1750, mounted where it measures desk illumination rather than looking directly into the lamp. It reports light level in lux, allowing the controller to distinguish daylight, adequate artificial light, and an under-lit workspace. This is more precise than a simple light-dependent resistor and permits a configurable comfort threshold for reading or computer work.",
    ],
    157: [
        "The controller would combine occupancy and lux rather than acting on either value alone. When the PIR detects a user and desk illumination is below the chosen threshold, it would switch on or dim up the desk lamp through a safe relay or smart-light interface. If daylight is already sufficient, the lamp would remain off. When no movement is detected for a defined period, the system would send a warning and then switch the lamp off; a manual switch must always override the automation.",
        "The platform would store only time-stamped occupancy state, lux, lamp state, and energy use. These records can reveal how long lights remain on unnecessarily and help tune the threshold and timeout. The user receives appropriate lighting without repeated manual adjustment, reduces wasted electricity when leaving the room, and retains privacy because the design uses motion and light measurements rather than a camera.",
    ],
    165: [
        "The selected city is Kota Bharu, Kelantan, Malaysia. Low-lying neighbourhoods, roads, and drains near Sungai Kelantan can become unsafe during intense rainfall and the monsoon when river and drain levels rise. Floodwater disrupts travel and may leave residents or drivers entering an affected area before a warning reaches them. The Department of Irrigation and Drainage already monitors Kota Bharu stations such as Sungai Kelantan at Tambatan D’Raja, so additional local sensor nodes could increase coverage between existing stations.",
        "An IoT network could place non-contact ultrasonic water-level sensors above critical drains and flood-prone roads, supported by tipping-bucket rain gauges at nearby locations. Each node would report water depth, rate of rise, rainfall intensity, battery condition, and location. Combining several nearby readings would give JPS Kelantan and the local authority earlier and more reliable evidence than a single visual report, allowing targeted warnings, drain inspections, road closures, evacuation preparation, and siren activation.",
    ],
    175: [
        "LoRaWAN is a low-power wide-area technology designed for small, infrequent sensor messages. A gateway can cover a broad urban area, although buildings, terrain, antenna height, and interference reduce the practical range. Battery nodes can sleep for most of the time and wake briefly to transmit a water-level packet, supporting long service life. Its limitations are low data rate, duty-cycle constraints, and less predictable latency, so it is unsuitable for continuous high-resolution video.",
        "5G uses licensed cellular infrastructure and provides higher capacity, mobility support, and lower-latency service classes. Coverage depends on the operator and radio band: lower-frequency cells cover wider areas, while high-frequency cells provide capacity over shorter distances. A 5G modem normally consumes more energy than a sleeping LoRaWAN node because it performs more complex radio processing and network signalling, although 5G is preferable when a site needs video, frequent large uploads, or rapid closed-loop control.",
        "LoRaWAN is the more practical primary link for Kota Bharu’s distributed river, drain, and low-lying-road sensors because each node sends only a few bytes, many locations lack convenient mains power, and battery replacement beside waterways creates cost and safety problems. Gateways can be mounted at JPS or local-authority facilities and forward data through fibre or cellular backhaul. Critical sites should use acknowledged alarms, repeated transmissions, local sirens, and a cellular fallback because no single wireless link should be the only flood-safety control.",
    ],
}


BLANK_RANGES = [
    range(88, 95),
    range(98, 107),
    range(112, 121),
    range(124, 132),
    range(137, 147),
    range(150, 155),
    range(158, 160),
    range(166, 173),
    range(176, 183),
    range(186, 200),
]


SOURCES = [
    "3rd Generation Partnership Project. (2023). Ultra reliable and low latency communications. https://www.3gpp.org/technologies/urlcc-2022",
    "Department of Irrigation and Drainage Malaysia. (2026). River water-level data: Kelantan. https://publicinfobanjir.water.gov.my/aras-air/data-paras-air/?lang=en&state=KEL",
    "International Telecommunication Union. (2012). Recommendation ITU-T Y.4000/Y.2060: Overview of the Internet of things. https://www.itu.int/itu-t/recommendations/rec.aspx?rec=Y.2060",
    "LoRa Alliance. (n.d.). What is LoRaWAN? https://lora-alliance.org/resource_hub/what-is-lorawan/",
    "World Health Organization. (2026). Cold chain equipment and dry store temperature mapping tool. https://www.who.int/publications/m/item/cold-chain-equipment-and-dry-store-temperature-mapping-tool",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color="7F8C8D", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def style_body(paragraph, before=0, after=6):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)


def insert_paragraph_after(paragraph, text, *, bold=False, italic=False, color=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    style_body(new_para)
    return new_para


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def insert_evidence_box_after(document, paragraph, text, height_inches):
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.4)
    row = table.rows[0]
    row.height = Inches(height_inches)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cell = row.cells[0]
    cell.width = Inches(6.4)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "FFF2CC")
    set_cell_margins(cell, 180, 220, 180, 220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(156, 0, 6)
    paragraph._p.addnext(table._tbl)


def add_sources_before(end_paragraph):
    heading = end_paragraph.insert_paragraph_before("Selected Technical Sources")
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    for run in heading.runs:
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(12)
    for source in SOURCES:
        p = end_paragraph.insert_paragraph_before(source)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def main():
    document = Document(SOURCE)
    original_paragraphs = list(document.paragraphs)

    original_paragraphs[8].text = "NAME: Chan Jing Yi"
    original_paragraphs[9].text = "STUDENT ID: SUOL2500321"
    for index in (8, 9):
        run = original_paragraphs[index].runs[0]
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)

    mcq_answers = ["B", "D", "B", "C", "B", "C", "A", "B", "C", "B"]
    table = document.tables[1]
    for question_number in range(1, 6):
        left_answer = table.cell(question_number, 1)
        right_answer = table.cell(question_number, 3)
        for cell, answer in (
            (left_answer, mcq_answers[question_number - 1]),
            (right_answer, mcq_answers[question_number + 4]),
        ):
            cell.text = answer
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "E2F0D9")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(12)

    for anchor_index, paragraphs in ANSWERS.items():
        anchor = original_paragraphs[anchor_index]
        if not anchor.runs:
            anchor.add_run("Answer:")
        for run in anchor.runs:
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(11)
        cursor = anchor
        for answer_paragraph in paragraphs:
            cursor = insert_paragraph_after(cursor, answer_paragraph)

    selfie_anchor = original_paragraphs[136]
    insert_evidence_box_after(
        document,
        selfie_anchor,
        "STUDENT ACTION REQUIRED — Insert a real selfie showing Chan Jing Yi beside the bedroom study desk and its lamp or wall-switch panel. Both the student and selected area must be clearly visible. A room-only photograph receives zero marks.",
        2.8,
    )

    diagram_anchor = original_paragraphs[185]
    insert_evidence_box_after(
        document,
        diagram_anchor,
        "STUDENT ACTION REQUIRED — Hand-draw on paper: Water-level sensor + rain gauge near Sungai Kelantan and flood-prone drains (Physical Layer) → LoRaWAN node and gateway (Network Layer) → IoT platform, database, threshold and rising-water analysis (Middleware Layer) → JPS Kelantan/local-authority dashboard, public mobile alert, road warning sign and local siren (Application Layer). Write SUOL2500321 and your signature in one corner, photograph the complete sheet, and replace this box with the image.",
        3.2,
    )

    for blank_range in BLANK_RANGES:
        for index in blank_range:
            delete_paragraph(original_paragraphs[index])

    add_sources_before(original_paragraphs[200])

    for section in document.sections:
        header = section.header
        header_p = header.paragraphs[0]
        header_p.clear()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_p.add_run("Chan Jing Yi  |  SUOL2500321  |  BCL1123 Internet of Things")
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)
        add_bottom_border(header_p, color="A6A6A6", size="4")

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.clear()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run("Chan Jing Yi  |  SUOL2500321  |  Page ")
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)
        add_page_field(footer_p)

    core = document.core_properties
    core.title = "BCL1123 Internet of Things Test Answers"
    core.subject = "May–August 2026 Test"
    core.author = "Chan Jing Yi"
    core.keywords = "BCL1123, Internet of Things, SUOL2500321"

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
