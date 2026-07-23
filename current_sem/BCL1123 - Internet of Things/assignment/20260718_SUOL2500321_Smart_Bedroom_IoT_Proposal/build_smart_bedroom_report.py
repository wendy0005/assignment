from __future__ import annotations

import hashlib
import html
import re
import shutil
import subprocess
import zipfile
from pathlib import Path

import markdown
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image as RLImage,
    KeepTogether,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents


ROOT = Path(__file__).resolve().parent
SOURCE_MD = ROOT / "20260718_SUOL2500321_Smart_Bedroom_IoT_Proposal.md"
REFERENCE_DOCX = ROOT / "RENAMETOSTUDENTIDASSIGNMENTCOVERPAGE.docx"
OUT_DOCX = ROOT / "20260718_SUOL2500321_Smart_Bedroom_IoT_Proposal.docx"
OUT_PDF = ROOT / "20260718_SUOL2500321_Smart_Bedroom_IoT_Proposal.pdf"
OUT_HTML = ROOT / "20260718_SUOL2500321_Smart_Bedroom_IoT_Proposal.html"
ASSETS = ROOT / "assets"

TNR = "/System/Library/Fonts/Supplemental/Times New Roman.ttf"
TNR_BOLD = "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"
TNR_ITALIC = "/System/Library/Fonts/Supplemental/Times New Roman Italic.ttf"
ARIAL = "/System/Library/Fonts/Supplemental/Arial.ttf"
ARIAL_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"

INK = "#17324D"
BLUE = "#0B6E99"
TEAL = "#008C95"
LIGHT = "#EAF4F7"
PALE = "#F5F8FA"
RED = "#B42318"
AMBER = "#B7791F"
GREEN = "#14804A"
GRAY = "#5B6573"


def font(path: str, size: int):
    return ImageFont.truetype(path, size)


def rounded(draw, box, fill, outline=None, radius=24, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def centered(draw, box, text, fnt, fill=INK, spacing=6):
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
    y = y1 + ((y2 - y1) - sum(heights) - spacing * (len(lines) - 1)) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text(((x1 + x2 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + spacing


def arrow(draw, start, end, fill=BLUE, width=7):
    draw.line([start, end], fill=fill, width=width)
    import math

    ang = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 22
    for delta in (2.55, -2.55):
        p = (end[0] + length * math.cos(ang + delta), end[1] + length * math.sin(ang + delta))
        draw.line([end, p], fill=fill, width=width)


def make_photo_placeholder(path: Path, label: str, subtitle: str):
    im = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(im)
    d.rectangle((20, 20, 1580, 880), outline=BLUE, width=8)
    d.rectangle((60, 60, 1540, 840), fill="#F2F7F9", outline="#A7C8D6", width=4)
    d.ellipse((690, 245, 910, 465), outline=BLUE, width=12)
    d.rectangle((620, 350, 980, 650), outline=BLUE, width=12)
    d.polygon([(645, 620), (790, 475), (890, 575), (940, 525), (965, 620)], fill="#B8D9E5")
    centered(d, (120, 90, 1480, 210), label, font(ARIAL_BOLD, 56), BLUE)
    centered(d, (160, 700, 1440, 810), subtitle, font(ARIAL, 34), GRAY)
    im.save(path)


def make_hardware_board(path: Path):
    im = Image.new("RGB", (2000, 1200), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 50), "SentinelSleep Hardware Selection", font=font(ARIAL_BOLD, 58), fill=INK)
    d.text((72, 120), "Wokwi-supported sensors, controller and actuators", font=font(ARIAL, 32), fill=GRAY)
    cards = [
        ("ESP32", "Controller + Wi-Fi", "chip"),
        ("DHT22", "Temperature + humidity", "thermo"),
        ("PIR", "Movement", "motion"),
        ("LDR", "Illumination", "sun"),
        ("MQ-2", "Supplementary gas", "gas"),
        ("Relays", "Light + fan", "relay"),
        ("Buzzer / RGB", "Local alert", "alert"),
        ("Servo", "Optional curtain", "servo"),
    ]
    for idx, (name, desc, kind) in enumerate(cards):
        col, row = idx % 4, idx // 4
        x, y = 65 + col * 485, 220 + row * 470
        rounded(d, (x, y, x + 430, y + 400), "#F7FAFC", "#B7CDD8", 26, 4)
        cx, cy = x + 215, y + 145
        if kind == "chip":
            d.rectangle((cx - 75, cy - 60, cx + 75, cy + 60), fill="#1B5E76", outline=INK, width=5)
            for off in range(-55, 76, 26):
                d.line((cx - 95, cy + off, cx - 75, cy + off), fill=INK, width=5)
                d.line((cx + 75, cy + off, cx + 95, cy + off), fill=INK, width=5)
            centered(d, (cx - 70, cy - 30, cx + 70, cy + 30), "ESP32", font(ARIAL_BOLD, 24), "white")
        elif kind == "thermo":
            d.ellipse((cx - 35, cy + 25, cx + 35, cy + 95), fill=RED)
            d.rounded_rectangle((cx - 14, cy - 80, cx + 14, cy + 55), radius=12, fill=RED)
            d.arc((cx + 50, cy - 65, cx + 150, cy + 35), 200, 340, fill=TEAL, width=8)
        elif kind == "motion":
            d.ellipse((cx - 30, cy - 85, cx + 30, cy - 25), fill=BLUE)
            d.line((cx, cy - 25, cx, cy + 70), fill=BLUE, width=12)
            d.line((cx, cy + 10, cx - 65, cy + 45), fill=BLUE, width=10)
            d.line((cx, cy + 10, cx + 65, cy + 45), fill=BLUE, width=10)
            d.arc((cx + 40, cy - 70, cx + 180, cy + 70), 220, 140, fill=TEAL, width=8)
        elif kind == "sun":
            d.ellipse((cx - 55, cy - 55, cx + 55, cy + 55), fill=AMBER)
            for a in range(0, 360, 45):
                import math
                p1 = (cx + 80 * math.cos(math.radians(a)), cy + 80 * math.sin(math.radians(a)))
                p2 = (cx + 120 * math.cos(math.radians(a)), cy + 120 * math.sin(math.radians(a)))
                d.line((p1, p2), fill=AMBER, width=8)
        elif kind == "gas":
            for off in (-55, 0, 55):
                d.arc((cx - 95 + off, cy - 40, cx + 25 + off, cy + 80), 200, 350, fill=RED, width=10)
        elif kind == "relay":
            d.rectangle((cx - 120, cy - 75, cx + 120, cy + 75), fill="#5DAE8B", outline=INK, width=5)
            d.rectangle((cx - 90, cy - 35, cx - 15, cy + 35), fill="#E8F4ED", outline=INK, width=3)
            d.ellipse((cx + 35, cy - 16, cx + 67, cy + 16), fill="#F2D16B", outline=INK)
        elif kind == "alert":
            d.polygon([(cx, cy - 100), (cx - 115, cy + 90), (cx + 115, cy + 90)], fill="#FDE7E4", outline=RED)
            d.text((cx - 12, cy - 45), "!", font=font(ARIAL_BOLD, 88), fill=RED)
        else:
            d.rectangle((cx - 110, cy - 45, cx + 50, cy + 45), fill="#5D89B3", outline=INK, width=4)
            d.line((cx + 50, cy, cx + 125, cy - 85), fill=INK, width=12)
            d.ellipse((cx + 108, cy - 105, cx + 142, cy - 71), fill=AMBER)
        centered(d, (x + 20, y + 270, x + 410, y + 325), name, font(ARIAL_BOLD, 36), INK)
        centered(d, (x + 20, y + 325, x + 410, y + 385), desc, font(ARIAL, 25), GRAY)
    im.save(path)


def make_architecture(path: Path):
    im = Image.new("RGB", (2400, 1250), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 45), "SentinelSleep Four-Layer System Architecture", font=font(ARIAL_BOLD, 58), fill=INK)
    columns = [
        ("EDGE", "#E9F5F7", ["DHT22", "PIR", "Photoresistor", "MQ-2", "ESP32 rules", "Relays / buzzer"]),
        ("CONNECTIVITY", "#EAF0FA", ["2.4 GHz Wi-Fi", "MQTT over TLS", "QoS + Last Will"]),
        ("CLOUD", "#F1ECFA", ["IoT broker", "Rules engine", "Time-series data", "Device state", "Notifications"]),
        ("APPLICATION", "#FFF4E4", ["Safety banner", "Live dashboard", "Modes + controls", "History + alerts"]),
    ]
    for i, (title, fill, items) in enumerate(columns):
        x1 = 70 + i * 580
        rounded(d, (x1, 180, x1 + 490, 1080), fill, "#8AA8B6", 30, 4)
        d.rectangle((x1, 180, x1 + 490, 280), fill=INK)
        centered(d, (x1, 180, x1 + 490, 280), title, font(ARIAL_BOLD, 34), "white")
        for j, item in enumerate(items):
            y = 330 + j * 120
            rounded(d, (x1 + 40, y, x1 + 450, y + 78), "white", "#AEC6D1", 18, 3)
            centered(d, (x1 + 50, y, x1 + 440, y + 78), item, font(ARIAL, 27), INK)
        if i < 3:
            arrow(d, (x1 + 500, 625), (x1 + 570, 625), BLUE, 8)
            arrow(d, (x1 + 570, 680), (x1 + 500, 680), TEAL, 8)
    d.text((760, 1140), "Telemetry and confirmed state  ->     <-  Commands and desired state", font=font(ARIAL_BOLD, 31), fill=GRAY)
    im.save(path)


def make_flowchart(path: Path):
    im = Image.new("RGB", (2100, 1500), "white")
    d = ImageDraw.Draw(im)
    d.text((65, 35), "Priority-Based Local Control Flow", font=font(ARIAL_BOLD, 56), fill=INK)
    boxes = {
        "read": (720, 150, 1380, 245, "Read and validate all sensors", LIGHT),
        "valid": (800, 315, 1300, 425, "Readings valid?", "#FFF4E4"),
        "fault": (70, 315, 610, 425, "Hold safe state\nLog fault and retry", "#FDE7E4"),
        "gas": (800, 510, 1300, 620, "Gas threshold active?", "#FFF4E4"),
        "alert": (1450, 500, 2030, 650, "Buzzer + red status\nUrgent MQTT alert\nEvacuation instruction", "#FDE7E4"),
        "night": (800, 705, 1300, 815, "Night hours?", "#FFF4E4"),
        "nightlight": (90, 895, 630, 1020, "Motion + very dark\nLow-output night light", "#EAF0FA"),
        "mainlight": (780, 895, 1320, 1020, "Motion + dark\nMain light / vacancy timer", "#EAF0FA"),
        "fan": (1470, 895, 2010, 1020, "Occupied + hot\nFan with hysteresis", "#EAF0FA"),
        "publish": (650, 1190, 1450, 1300, "Update local display, publish state, wait 2 seconds", "#E9F5F7"),
    }
    for key, (x1, y1, x2, y2, text, fill) in boxes.items():
        rounded(d, (x1, y1, x2, y2), fill, RED if key in ("fault", "alert") else "#83A8B8", 22, 4)
        centered(d, (x1 + 15, y1 + 5, x2 - 15, y2 - 5), text, font(ARIAL_BOLD if key in ("gas", "valid", "night") else ARIAL, 29), INK)
    arrow(d, (1050, 245), (1050, 315))
    arrow(d, (800, 370), (610, 370), RED)
    d.text((660, 330), "No", font=font(ARIAL_BOLD, 24), fill=RED)
    arrow(d, (1050, 425), (1050, 510))
    d.text((1070, 455), "Yes", font=font(ARIAL_BOLD, 24), fill=GREEN)
    arrow(d, (1300, 565), (1450, 565), RED)
    d.text((1350, 525), "Yes", font=font(ARIAL_BOLD, 24), fill=RED)
    arrow(d, (1050, 620), (1050, 705))
    d.text((1070, 650), "No", font=font(ARIAL_BOLD, 24), fill=GREEN)
    arrow(d, (800, 760), (630, 945))
    d.text((690, 790), "Yes", font=font(ARIAL_BOLD, 24), fill=BLUE)
    arrow(d, (1050, 815), (1050, 895))
    d.text((1070, 835), "No", font=font(ARIAL_BOLD, 24), fill=BLUE)
    arrow(d, (1300, 760), (1740, 895))
    d.text((1390, 790), "Comfort path", font=font(ARIAL_BOLD, 24), fill=BLUE)
    for x in (360, 1050, 1740):
        arrow(d, (x, 1020), (1050, 1190), TEAL)
    arrow(d, (650, 1245), (250, 1245), TEAL)
    d.line((250, 1245, 250, 200, 720, 200), fill=TEAL, width=7)
    arrow(d, (250, 200), (720, 200), TEAL)
    im.save(path)


def make_dashboard(path: Path):
    im = Image.new("RGB", (1900, 1300), "#E9EEF2")
    d = ImageDraw.Draw(im)
    rounded(d, (140, 55, 1760, 1245), "white", "#AEBEC8", 36, 5)
    d.rectangle((140, 55, 1760, 180), fill=INK)
    d.text((205, 88), "SentinelSleep", font=font(ARIAL_BOLD, 44), fill="white")
    d.text((1370, 102), "AUTO MODE", font=font(ARIAL_BOLD, 26), fill="#9FE1D1")
    rounded(d, (200, 220, 1700, 340), "#E8F6EF", GREEN, 22, 4)
    d.ellipse((240, 255, 290, 305), fill=GREEN)
    d.text((325, 248), "ROOM SAFE", font=font(ARIAL_BOLD, 38), fill=GREEN)
    d.text((1110, 255), "Last update: 8:30 PM", font=font(ARIAL, 25), fill=GRAY)
    cards = [("29.1 C", "Temperature", AMBER), ("68%", "Humidity", BLUE), ("42 lux", "Light", TEAL), ("Detected", "Occupancy", GREEN)]
    for i, (value, label, color) in enumerate(cards):
        x = 200 + i * 375
        rounded(d, (x, 390, x + 330, 565), "#F8FAFC", "#C7D5DD", 18, 3)
        d.text((x + 28, 420), value, font=font(ARIAL_BOLD, 38), fill=color)
        d.text((x + 28, 500), label, font=font(ARIAL, 24), fill=GRAY)
    d.text((205, 625), "DEVICES", font=font(ARIAL_BOLD, 28), fill=INK)
    devices = [("Main light", "ON", "Occupied + dark"), ("Fan", "ON", "Occupied + 29.1 C"), ("Night light", "OFF", "Outside night hours")]
    for i, (name, state, why) in enumerate(devices):
        y = 680 + i * 115
        rounded(d, (200, y, 1100, y + 88), "#F7FAFC", "#CBD8DF", 16, 2)
        d.text((235, y + 22), name, font=font(ARIAL_BOLD, 27), fill=INK)
        d.text((500, y + 24), state, font=font(ARIAL_BOLD, 25), fill=GREEN if state == "ON" else GRAY)
        d.text((650, y + 24), why, font=font(ARIAL, 23), fill=GRAY)
    d.text((1180, 625), "MODES", font=font(ARIAL_BOLD, 28), fill=INK)
    for i, mode in enumerate(("Auto", "Sleep", "Study", "Away")):
        y = 680 + i * 92
        rounded(d, (1180, y, 1600, y + 65), BLUE if i == 0 else "#F7FAFC", "#AFC5D0", 18, 2)
        centered(d, (1180, y, 1600, y + 65), mode, font(ARIAL_BOLD, 24), "white" if i == 0 else INK)
    rounded(d, (1180, 1080, 1600, 1155), "#FDE7E4", RED, 18, 3)
    centered(d, (1180, 1080, 1600, 1155), "View alert history", font(ARIAL_BOLD, 23), RED)
    im.save(path)


def render_mermaid(source: Path, png_output: Path, svg_output: Path):
    """Render one Mermaid source to reproducible static assets."""
    command = [
        "npx",
        "--yes",
        "@mermaid-js/mermaid-cli@11.16.0",
        "-i",
        str(source),
        "-b",
        "white",
        "-w",
        "2400",
    ]
    subprocess.run([*command, "-o", str(svg_output)], check=True)
    subprocess.run([*command, "-o", str(png_output), "-s", "2"], check=True)


def build_assets():
    ASSETS.mkdir(exist_ok=True)
    make_photo_placeholder(ASSETS / "figure1_bedroom_photo_placeholder.png", "ORIGINAL BEDROOM PHOTO REQUIRED", "Replace this panel with a wide photograph of the selected bedroom.")
    make_photo_placeholder(ASSETS / "figure2_controls_photo_placeholder.png", "ORIGINAL CONTROL-POINT PHOTO REQUIRED", "Replace this panel with the existing light/fan controls or proposed sensor location.")
    make_hardware_board(ASSETS / "figure3_hardware_selection.png")
    render_mermaid(
        ASSETS / "figure4_architecture.mmd",
        ASSETS / "figure4_architecture.png",
        ASSETS / "figure4_architecture.svg",
    )
    render_mermaid(
        ASSETS / "figure5_control_flow.mmd",
        ASSETS / "figure5_control_flow.png",
        ASSETS / "figure5_control_flow.svg",
    )
    make_dashboard(ASSETS / "figure6_dashboard_wireframe.png")


def iter_all_paragraphs(parent):
    for p in getattr(parent, "paragraphs", []):
        yield p
    for t in getattr(parent, "tables", []):
        for row in t.rows:
            for cell in row.cells:
                yield from iter_all_paragraphs(cell)


def set_run_font(run, size=12, bold=None, italic=None, color=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def fill_cover(doc: Document):
    name_done = id_done = False
    for p in iter_all_paragraphs(doc):
        txt = p.text.strip()
        if txt.startswith("Module Code and Title") and "BCL1123" not in txt:
            r = p.add_run(" BCL1123 - Internet of Things")
            set_run_font(r, bold=True)
        elif txt.startswith("Program of Study") and "Degree IoT ODL" not in txt:
            r = p.add_run(" Degree IoT ODL")
            set_run_font(r, bold=True)
        elif txt.startswith("Title of Assignment") and "Proposal Report" not in txt:
            r = p.add_run(" Proposal Report & Video")
            set_run_font(r, bold=True)
        elif txt.startswith("Type (* please tick") and "[X]" not in txt:
            r = p.add_run("   [X] INDIVIDUAL")
            set_run_font(r, bold=True)
        elif txt.startswith("Name:") and not name_done:
            r = p.add_run(" Chan Jing Yi")
            set_run_font(r, bold=True)
            name_done = True
        elif txt.startswith("ID:") and not id_done:
            r = p.add_run(" SUOL2500321")
            set_run_font(r, bold=True)
            id_done = True


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = int(sum(widths_inches) * 1440)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[idx])
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(int(widths_inches[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_section(section, orientation="portrait", title="BCL1123 Internet of Things | SentinelSleep Proposal", restart=False):
    section.orientation = WD_ORIENT.LANDSCAPE if orientation == "landscape" else WD_ORIENT.PORTRAIT
    if orientation == "landscape":
        section.page_width, section.page_height = Mm(297), Mm(210)
    else:
        section.page_width, section.page_height = Mm(210), Mm(297)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    hp = section.header.paragraphs[0]
    hp.text = title
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs:
        set_run_font(r, 9, bold=True, color=GRAY)
    fp = section.footer.paragraphs[0]
    fp.text = "Chan Jing Yi | SUOL2500321 | "
    add_page_field(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in fp.runs:
        set_run_font(r, 9, color=GRAY)
    if restart:
        sect_pr = section._sectPr
        pg_num = sect_pr.find(qn("w:pgNumType"))
        if pg_num is None:
            pg_num = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num)
        pg_num.set(qn("w:start"), "1")


def add_toc_field(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Table of contents updates automatically when opened in Microsoft Word."
    text_run.append(text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text_run, fld_end])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Inches(0.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    for name, size, before, after in (("Heading 1", 16, 14, 8), ("Heading 2", 14, 12, 6), ("Heading 3", 12, 8, 4)):
        try:
            s = styles[name]
        except KeyError:
            s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            s.base_style = normal
        s.font.name = "Times New Roman"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(INK.replace("#", ""))
        s.paragraph_format.first_line_indent = Inches(0)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True


def clean_md(text):
    text = text.replace("**", "").replace("`", "")
    text = re.sub(r"(?<!\w)\*(.*?)\*", r"\1", text)
    text = text.replace("  ", " ")
    return text.strip()


def add_docx_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = False
    r = p.add_run(clean_md(text))
    set_run_font(r, 12, italic=True)


def add_docx_image(doc, image_path, width_inches):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    drawing = run._r.find(qn("w:drawing"))
    if drawing is not None:
        doc_pr = drawing.find(".//" + qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", image_path.stem.replace("_", " "))
            doc_pr.set("title", image_path.stem)


def add_docx_table(doc, rows, landscape_mode=False):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    total_width = 9.5 if landscape_mode else 6.5
    if len(rows[0]) == 2:
        widths = [total_width * 0.34, total_width * 0.66]
    elif len(rows[0]) == 3:
        widths = [total_width * 0.22, total_width * 0.35, total_width * 0.43]
    elif len(rows[0]) == 4:
        widths = [total_width * 0.34] + [total_width * 0.22] * 3
    else:
        widths = [total_width / len(rows[0])] * len(rows[0])
    set_table_width(table, widths)
    set_repeat_table_header(table.rows[0])
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = clean_md(value)
            if ri == 0:
                set_cell_shading(cell, "DDEBF2")
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(2)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 and len(value) < 30 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    set_run_font(r, 10.5, bold=(ri == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_docx_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    # Minimal inline markdown handling.
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        italic = not bold and part.startswith("*") and part.endswith("*")
        code = part.startswith("`") and part.endswith("`")
        content = part[2:-2] if bold else part[1:-1] if (italic or code) else part
        r = p.add_run(content)
        set_run_font(r, 12, bold=bold, italic=italic)
    return p


def parse_markdown_lines():
    return SOURCE_MD.read_text(encoding="utf-8").splitlines()


def build_docx():
    expected_hash = "0dbe6fbcbe9c0e811ba1c8fb59fe4312bd4502642016cc1486cbd4193437ed74"
    actual = hashlib.sha256(REFERENCE_DOCX.read_bytes()).hexdigest()
    if actual != expected_hash:
        raise RuntimeError("Cover template changed; fresh distillation is required")
    work = ROOT / ".smart_bedroom_working.docx"
    shutil.copy2(REFERENCE_DOCX, work)
    doc = Document(work)
    fill_cover(doc)
    configure_styles(doc)
    cover = doc.sections[0]
    cover.different_first_page_header_footer = False
    # Preserve official cover header and add report section.
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(sec, "portrait", restart=True)
    current_landscape = False
    started = False
    lines = parse_markdown_lines()
    i = 0
    mermaid_index = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line == "## Table of Contents":
            started = True
            p = doc.add_paragraph("Table of Contents", style="Heading 1")
            p.paragraph_format.page_break_before = False
            add_toc_field(doc)
            doc.add_page_break()
            i += 1
            while i < len(lines) and (not lines[i].startswith("## 1.")):
                i += 1
            continue
        if not started:
            i += 1
            continue
        if line.startswith("### 3.4 System architecture"):
            sec = doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(sec, "landscape", "BCL1123 | SentinelSleep System Design")
            current_landscape = True
        elif line.startswith("### 3.6 Dashboard UI prototype") and current_landscape:
            sec = doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(sec, "portrait", "BCL1123 | SentinelSleep Proposal")
            current_landscape = False
        if line.startswith("## "):
            if not line.startswith("## 1."):
                doc.add_page_break()
            doc.add_paragraph(clean_md(line[3:]), style="Heading 1")
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(clean_md(line[4:]), style="Heading 2")
            i += 1
            continue
        if line.startswith("```mermaid"):
            mermaid_index += 1
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                i += 1
            asset = ASSETS / ("figure4_architecture.png" if mermaid_index == 1 else "figure5_control_flow.png")
            width = 9.2 if mermaid_index == 1 and current_landscape else 6.6 if mermaid_index == 2 else 6.2
            add_docx_image(doc, asset, width)
            i += 1
            continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            add_docx_table(doc, rows, current_landscape)
            continue
        if line.startswith("**Figure 1 placeholder"):
            add_docx_image(doc, ASSETS / "figure1_bedroom_photo_placeholder.png", 6.1)
            if i + 2 < len(lines):
                add_docx_caption(doc, clean_md(lines[i + 2]))
                i += 3
            else:
                i += 1
            continue
        if line.startswith("**Figure 2 placeholder"):
            add_docx_image(doc, ASSETS / "figure2_controls_photo_placeholder.png", 6.1)
            if i + 2 < len(lines):
                add_docx_caption(doc, clean_md(lines[i + 2]))
                i += 3
            else:
                i += 1
            continue
        if line.startswith("**Figure 3."):
            add_docx_image(doc, ASSETS / "figure3_hardware_selection.png", 5.35)
            add_docx_caption(doc, line)
            i += 1
            continue
        if line.startswith("**Figure 4.") or line.startswith("**Figure 5."):
            add_docx_caption(doc, line)
            i += 1
            continue
        if line.startswith("**Figure 6."):
            add_docx_image(doc, ASSETS / "figure6_dashboard_wireframe.png", 6.2)
            add_docx_caption(doc, line)
            i += 1
            continue
        if line.startswith("**Shared video URL:"):
            p = add_docx_para(doc, line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(0)
            i += 1
            continue
        if line.startswith("> "):
            p = add_docx_para(doc, line[2:])
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.keep_together = True
            for r in p.runs:
                set_run_font(r, 12, bold=True, color=RED)
            i += 1
            continue
        if line.strip() and not re.match(r"^\d+\. ", line):
            add_docx_para(doc, line)
        i += 1
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    doc.core_properties.title = "SentinelSleep Smart Bedroom IoT Proposal"
    doc.core_properties.author = "Chan Jing Yi"
    doc.core_properties.subject = "BCL1123 Internet of Things - Proposal Report"
    doc.core_properties.keywords = "IoT, ESP32, smart bedroom, Wokwi, MQTT"
    doc.save(OUT_DOCX)
    patch_decorative_alt_text(OUT_DOCX)
    work.unlink(missing_ok=True)


def patch_decorative_alt_text(docx_path: Path):
    temp = docx_path.with_suffix(".altfix.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename in ("word/document.xml", "word/header1.xml"):
                from lxml import etree

                root = etree.fromstring(data)
                ns = {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}
                for node in root.xpath("//wp:docPr", namespaces=ns):
                    if not node.get("descr"):
                        if info.filename == "word/header1.xml":
                            node.set("descr", "Official faculty logo or decorative cover-page header element")
                            node.set("title", "Official faculty cover header")
                        else:
                            node.set("descr", "Decorative official cover-page layout element")
                            node.set("title", "Official cover-page element")
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(info, data)
    temp.replace(docx_path)


def md_inline(text):
    text = html.escape(text)
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)
    text = re.sub(r"`(.*?)`", r"<font name='Courier'>\1</font>", text)
    return text


class ReportDocTemplate(BaseDocTemplate):
    def __init__(self, filename, **kw):
        super().__init__(filename, **kw)
        self._heading_seq = 0

    def beforeDocument(self):
        super().beforeDocument()
        self._heading_seq = 0

    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph):
            style = flowable.style.name
            if style in ("ReportH1", "ReportH2"):
                level = 0 if style == "ReportH1" else 1
                text = flowable.getPlainText()
                self._heading_seq += 1
                key = f"h{self._heading_seq}"
                self.canv.bookmarkPage(key)
                self.canv.addOutlineEntry(text, key, level=level, closed=False)
                if level == 0:
                    self.notify("TOCEntry", (level, text, self.page, key))


def pdf_header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Times-Roman", 9)
    canvas.setFillColor(colors.HexColor(GRAY))
    width, height = canvas._pagesize
    canvas.drawCentredString(width / 2, height - 28, "BCL1123 Internet of Things | SentinelSleep Proposal")
    canvas.drawCentredString(width / 2, 25, f"Chan Jing Yi | SUOL2500321 | Page {doc.page}")
    canvas.restoreState()


def cover_page(canvas, doc):
    canvas.saveState()
    width, height = A4
    logo = ROOT / ".cover_logo.png"
    if not logo.exists():
        import zipfile
        with zipfile.ZipFile(REFERENCE_DOCX) as z:
            logo.write_bytes(z.read("word/media/image1.png"))
    canvas.drawImage(str(logo), width / 2 - 65, height - 95, width=130, height=60, preserveAspectRatio=True, mask="auto")
    canvas.setFont("Times-Bold", 12)
    canvas.setFillColor(colors.HexColor(INK))
    canvas.drawCentredString(width / 2, height - 120, "Faculty of Engineering, Built Environment, and Information Technology")
    canvas.setStrokeColor(colors.HexColor(BLUE))
    canvas.setLineWidth(2)
    canvas.line(65, height - 140, width - 65, height - 140)
    y = height - 190
    canvas.setFont("Times-Bold", 17)
    canvas.drawCentredString(width / 2, y, "PROPOSAL REPORT & VIDEO")
    canvas.setFont("Times-Bold", 20)
    canvas.setFillColor(colors.HexColor(BLUE))
    canvas.drawCentredString(width / 2, y - 52, "SentinelSleep")
    canvas.setFont("Times-Roman", 13)
    canvas.setFillColor(colors.black)
    canvas.drawCentredString(width / 2, y - 78, "An Occupancy-Aware Smart Bedroom Comfort, Energy and Safety System")
    info = [
        ("Module", "BCL1123 - Internet of Things"),
        ("Program", "Degree IoT ODL"),
        ("Lecturer", "Lee Thian Seng"),
        ("Semester", "May-August 2026"),
        ("Student", "Chan Jing Yi"),
        ("Student ID", "SUOL2500321"),
        ("Submission date", "As in LMS"),
        ("Assessment type", "Individual"),
    ]
    y2 = y - 150
    for label, value in info:
        canvas.setFont("Times-Bold", 12)
        canvas.drawString(120, y2, f"{label}:")
        canvas.setFont("Times-Roman", 12)
        canvas.drawString(235, y2, value)
        y2 -= 27
    canvas.setFont("Times-Italic", 10.5)
    text = canvas.beginText(85, 145)
    text.setLeading(14)
    text.textLine("Declaration: I declare that this assignment is my own work and acknowledge the University's")
    text.textLine("academic-integrity requirements. Signature: ______________________________")
    canvas.drawText(text)
    canvas.setFont("Times-Bold", 9.5)
    canvas.setFillColor(colors.HexColor(RED))
    canvas.drawCentredString(width / 2, 90, "Before submission: insert original bedroom photographs and the accessible video URL.")
    canvas.restoreState()


def pdf_styles():
    pdfmetrics.registerFont(TTFont("Times-Roman", TNR))
    pdfmetrics.registerFont(TTFont("Times-Bold", TNR_BOLD))
    pdfmetrics.registerFont(TTFont("Times-Italic", TNR_ITALIC))
    pdfmetrics.registerFontFamily("Times-Roman", normal="Times-Roman", bold="Times-Bold", italic="Times-Italic")
    return {
        "body": ParagraphStyle("Body", fontName="Times-Roman", fontSize=12, leading=18, alignment=TA_JUSTIFY, firstLineIndent=0.5 * inch, spaceAfter=7, splitLongWords=True),
        "h1": ParagraphStyle("ReportH1", fontName="Times-Bold", fontSize=16, leading=20, textColor=colors.HexColor(INK), spaceBefore=10, spaceAfter=9, keepWithNext=True),
        "h2": ParagraphStyle("ReportH2", fontName="Times-Bold", fontSize=14, leading=18, textColor=colors.HexColor(INK), spaceBefore=9, spaceAfter=7, keepWithNext=True),
        "caption": ParagraphStyle("Caption", fontName="Times-Italic", fontSize=10.5, leading=14, alignment=TA_CENTER, spaceAfter=9),
        "toc": ParagraphStyle("TOC", fontName="Times-Roman", fontSize=12, leading=18, leftIndent=12, firstLineIndent=0, spaceAfter=4),
        "note": ParagraphStyle("Note", fontName="Times-Bold", fontSize=11.5, leading=16, textColor=colors.HexColor(RED), backColor=colors.HexColor("#FDE7E4"), borderPadding=8, spaceAfter=10),
        "table": ParagraphStyle("TableText", fontName="Times-Roman", fontSize=9.5, leading=12, alignment=TA_LEFT),
        "table_h": ParagraphStyle("TableHead", fontName="Times-Bold", fontSize=9.5, leading=12, alignment=TA_CENTER),
    }


def build_pdf():
    st = pdf_styles()
    portrait_frame = Frame(0.75 * inch, 0.65 * inch, A4[0] - 1.5 * inch, A4[1] - 1.3 * inch, id="portrait")
    land = landscape(A4)
    landscape_frame = Frame(0.65 * inch, 0.65 * inch, land[0] - 1.3 * inch, land[1] - 1.3 * inch, id="landscape")
    cover_frame = Frame(0, 0, A4[0], A4[1], leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0, id="cover")
    doc = ReportDocTemplate(str(OUT_PDF), pagesize=A4, rightMargin=0.75 * inch, leftMargin=0.75 * inch, topMargin=0.65 * inch, bottomMargin=0.65 * inch, title="SentinelSleep Smart Bedroom IoT Proposal", author="Chan Jing Yi")
    doc.addPageTemplates([
        PageTemplate(id="cover", pagesize=A4, frames=[cover_frame], onPage=cover_page),
        PageTemplate(id="portrait", pagesize=A4, frames=[portrait_frame], onPage=pdf_header_footer),
        PageTemplate(id="landscape", pagesize=land, frames=[landscape_frame], onPage=pdf_header_footer),
    ])
    story = [Spacer(1, A4[1] - 1), NextPageTemplate("portrait"), PageBreak()]
    story.append(Paragraph("Table of Contents", st["h1"]))
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("TOC1", fontName="Times-Roman", fontSize=11, leading=15, leftIndent=12, firstLineIndent=-6, spaceBefore=4),
        ParagraphStyle("TOC2", fontName="Times-Roman", fontSize=10, leading=13, leftIndent=28, firstLineIndent=-6, spaceBefore=2),
    ]
    story.extend([toc, PageBreak()])
    lines = parse_markdown_lines()
    started = False
    i = 0
    mermaid_index = 0
    current_landscape = False
    while i < len(lines):
        line = lines[i].rstrip()
        if line == "## 1. Executive Summary":
            started = True
        if not started:
            i += 1
            continue
        if line.startswith("### 3.4 System architecture"):
            story.extend([NextPageTemplate("landscape"), PageBreak()])
            current_landscape = True
        elif line.startswith("### 3.6 Dashboard UI prototype") and current_landscape:
            story.extend([NextPageTemplate("portrait"), PageBreak()])
            current_landscape = False
        if line.startswith("## "):
            if line != "## 1. Executive Summary":
                story.append(PageBreak())
            story.append(Paragraph(md_inline(line[3:]), st["h1"]))
            i += 1
            continue
        if line.startswith("### "):
            story.append(Paragraph(md_inline(line[4:]), st["h2"]))
            i += 1
            continue
        if line.startswith("```mermaid"):
            mermaid_index += 1
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                i += 1
            asset = ASSETS / ("figure4_architecture.png" if mermaid_index == 1 else "figure5_control_flow.png")
            max_w = 9.6 * inch if current_landscape else 6.5 * inch
            img = RLImage(str(asset))
            max_h = 4.1 * inch if mermaid_index == 2 else 5.4 * inch
            scale = min(max_w / img.imageWidth, max_h / img.imageHeight)
            img.drawWidth, img.drawHeight = img.imageWidth * scale, img.imageHeight * scale
            story.append(img)
            i += 1
            continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            data = [[Paragraph(md_inline(c), st["table_h"] if ri == 0 else st["table"]) for c in row] for ri, row in enumerate(rows)]
            avail = 9.6 * inch if current_landscape else 6.5 * inch
            n = len(rows[0])
            widths = [avail / n] * n
            if n == 3:
                widths = [avail * 0.22, avail * 0.35, avail * 0.43]
            elif n == 4:
                widths = [avail * 0.34, avail * 0.22, avail * 0.22, avail * 0.22]
            tbl = Table(data, colWidths=widths, repeatRows=1, hAlign="CENTER")
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DDEBF2")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#8AA8B6")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.extend([tbl, Spacer(1, 8)])
            continue
        if line.startswith("**Figure 1 placeholder"):
            img = RLImage(str(ASSETS / "figure1_bedroom_photo_placeholder.png"), width=6.2 * inch, height=3.49 * inch)
            caption = clean_md(lines[i + 2]) if i + 2 < len(lines) else "Figure 1. Original bedroom photograph placeholder."
            story.append(KeepTogether([img, Paragraph(md_inline(caption), st["caption"])]))
            i += 3
            continue
        if line.startswith("**Figure 2 placeholder"):
            img = RLImage(str(ASSETS / "figure2_controls_photo_placeholder.png"), width=6.2 * inch, height=3.49 * inch)
            caption = clean_md(lines[i + 2]) if i + 2 < len(lines) else "Figure 2. Original control-point photograph placeholder."
            story.append(KeepTogether([img, Paragraph(md_inline(caption), st["caption"])]))
            i += 3
            continue
        if line.startswith("**Figure 3."):
            img = RLImage(str(ASSETS / "figure3_hardware_selection.png"), width=5.35 * inch, height=3.21 * inch)
            story.append(KeepTogether([img, Paragraph(md_inline(clean_md(line)), st["caption"])]))
            i += 1
            continue
        if line.startswith("**Figure 4.") or line.startswith("**Figure 5."):
            story.append(Paragraph(md_inline(clean_md(line)), st["caption"]))
            i += 1
            continue
        if line.startswith("**Figure 6."):
            img = RLImage(str(ASSETS / "figure6_dashboard_wireframe.png"), width=6.15 * inch, height=4.21 * inch)
            story.append(KeepTogether([img, Paragraph(md_inline(clean_md(line)), st["caption"])]))
            i += 1
            continue
        if line.startswith("**Shared video URL:"):
            special = ParagraphStyle("VideoLink", parent=st["body"], alignment=TA_LEFT, firstLineIndent=0, wordWrap="CJK")
            story.append(Paragraph(md_inline(line), special))
            i += 1
            continue
        if line.startswith("> "):
            story.append(Paragraph(md_inline(line[2:]), st["note"]))
            i += 1
            continue
        if line.strip() and not re.match(r"^\d+\. ", line):
            story.append(Paragraph(md_inline(line), st["body"]))
        i += 1
    doc.multiBuild(story)
    (ROOT / ".cover_logo.png").unlink(missing_ok=True)


def build_html():
    body = markdown.markdown(
        SOURCE_MD.read_text(encoding="utf-8"),
        extensions=["tables", "fenced_code", "toc"],
        output_format="html5",
    )
    figure_inserts = {
        "<p><strong>Figure 1 placeholder": '<figure><img src="assets/figure1_bedroom_photo_placeholder.png" alt="Original bedroom photograph placeholder"></figure><p><strong>Figure 1 placeholder',
        "<p><strong>Figure 2 placeholder": '<figure><img src="assets/figure2_controls_photo_placeholder.png" alt="Original control-point photograph placeholder"></figure><p><strong>Figure 2 placeholder',
        "<p><strong>Figure 3.": '<figure><img src="assets/figure3_hardware_selection.png" alt="Illustrated selection of ESP32, sensors, and actuators"></figure><p><strong>Figure 3.',
        "<p><strong>Figure 6.": '<figure><img src="assets/figure6_dashboard_wireframe.png" alt="SentinelSleep mobile dashboard wireframe"></figure><p><strong>Figure 6.',
    }
    for marker, replacement in figure_inserts.items():
        body = body.replace(marker, replacement)
    mermaid_images = iter(
        [
            ("assets/figure4_architecture.png", "SentinelSleep four-layer system architecture"),
            ("assets/figure5_control_flow.png", "Priority-based SentinelSleep local control flow"),
        ]
    )

    def replace_mermaid_block(_match):
        source, alt = next(mermaid_images)
        return f'<figure><img src="{source}" alt="{alt}"></figure>'

    body = re.sub(
        r'<pre><code class="language-mermaid">.*?</code></pre>',
        replace_mermaid_block,
        body,
        flags=re.DOTALL,
    )
    page = f"""<!doctype html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>SentinelSleep Smart Bedroom IoT Proposal</title>
<style>
@page {{ size: A4; margin: 22mm 20mm 20mm; }}
body {{ max-width: 900px; margin: 0 auto; font-family: 'Times New Roman', serif; font-size: 12pt; line-height: 1.5; color: #172b3a; }}
h1 {{ text-align:center; color:{INK}; margin-top:2rem; }} h2 {{ color:{INK}; border-bottom:1px solid #b9ccd6; padding-bottom:.25rem; page-break-after:avoid; }} h3 {{ color:{BLUE}; page-break-after:avoid; }}
p {{ text-align:justify; text-indent:2em; }} blockquote {{ border-left:5px solid {RED}; background:#fde7e4; margin:1rem 0; padding:.7rem 1rem; }} blockquote p {{ text-indent:0; font-weight:bold; }}
table {{ width:100%; border-collapse:collapse; margin:1rem 0; font-size:10pt; }} th {{ background:#ddebf2; }} th,td {{ border:1px solid #8aa8b6; padding:7px; vertical-align:middle; }}
figure {{ margin:1rem auto .25rem; text-align:center; page-break-inside:avoid; }} figure img {{ max-width:100%; max-height:650px; object-fit:contain; }}
.mermaid {{ text-align:center; margin:1rem auto; page-break-inside:avoid; }} code {{ white-space:pre-wrap; }}
a {{ color:{BLUE}; overflow-wrap:anywhere; }}
@media print {{ h2 {{ page-break-before:always; }} h2:first-of-type {{ page-break-before:auto; }} body {{ max-width:none; }} }}
</style></head><body>{body}
</body></html>"""
    OUT_HTML.write_text(page, encoding="utf-8")


if __name__ == "__main__":
    build_assets()
    build_docx()
    build_pdf()
    build_html()
    print(OUT_DOCX)
    print(OUT_PDF)
    print(OUT_HTML)
